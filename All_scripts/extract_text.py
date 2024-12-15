import fitz  # PyMuPDF
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Chemin du fichier PDF spécifique
fichier_pdf = "C:\\Users\\Utilisateur\\Documents\\Devoirs et Cours 5A\\PTUT\\Sources RAG - Fine-tuning\\santexpo2024-dev\\pdf_all\\guide-des-carrieres-des-personnels-non-medicaux-de-la-fonction-publique-hospitaliere-juillet-2009.pdf"

# Chemin du dossier où tu veux enregistrer le fichier DOCX
dossier_sortie = "C:\\Users\\Utilisateur\\Documents\\Devoirs et Cours 5A\\PTUT\\Test PDF extract"

# Nom du fichier DOCX
nom_fichier = "resultat_extraction.docx"

# Combinaison du chemin complet pour le fichier de sortie
chemin_fichier_sortie = os.path.join(dossier_sortie, nom_fichier)

# Fonction d'extraction de texte
def extraire_texte_pdf(fichier_pdf):
    texte_complet = ""
    with fitz.open(fichier_pdf) as doc:
        for page_num, page in enumerate(doc):
            texte_complet += page.get_text()
    return texte_complet

# Fonction pour extraire les images
def extraire_images_pdf(fichier_pdf):
    images = []
    with fitz.open(fichier_pdf) as doc:
        for page_num, page in enumerate(doc):
            for img in page.get_images(full=True):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_data = base_image["image"]
                images.append(image_data)  # Ajoute l'image en bytes
    return images

# Fonction pour extraire les liens
def extraire_liens_pdf(fichier_pdf):
    liens = []
    with fitz.open(fichier_pdf) as doc:
        for page_num, page in enumerate(doc):
            for link in page.get_links():
                liens.append(link.get("uri"))  # Ajoute l'URL du lien
    return liens

# Fonction pour extraire les tableaux (approche basique)
def extraire_tableaux_pdf(fichier_pdf):
    tableaux = []
    with fitz.open(fichier_pdf) as doc:
        for page_num, page in enumerate(doc):
            text = page.get_text("dict")  # Extraire le texte au format dict
            for block in text["blocks"]:
                if block["type"] == 0:  # Bloc de texte
                    if "text" in block:
                        tableaux.append(block["text"])
    return tableaux

# Fonction pour créer un document Word et y insérer du texte, des images et des liens
def creer_document_word(texte, images, liens, tableaux, chemin_sortie):
    doc = Document()
    
    # Ajouter le texte extrait
    doc.add_paragraph(f"Texte extrait :\n\n{texte}")
    doc.add_paragraph("-" * 80)
    
    # Ajouter les liens extraits
    doc.add_paragraph("Liens extraits :")
    for lien in liens:
        doc.add_paragraph(lien)
    
    doc.add_paragraph("-" * 80)
    
    # Ajouter les tableaux extraits (approche basique)
    doc.add_paragraph("Tableaux extraits :")
    for tableau in tableaux:
        doc.add_paragraph(tableau)
    
    doc.add_paragraph("-" * 80)
    
    # Ajouter les images extraites
    doc.add_paragraph(f"Nombre d'images extraites : {len(images)}")
    for idx, image_data in enumerate(images):
        image_stream = BytesIO(image_data)  # Convertir l'image en stream
        doc.add_paragraph(f"Image {idx + 1}:")
        doc.add_picture(image_stream, width=Inches(4))  # Ajouter l'image avec une taille fixe
    
    # Sauvegarder le document Word
    doc.save(chemin_sortie)

# Extraction des données et création du fichier DOCX
# Extraction des données
texte = extraire_texte_pdf(fichier_pdf)
images = extraire_images_pdf(fichier_pdf)  # Liste d'images en bytes
liens = extraire_liens_pdf(fichier_pdf)
tableaux = extraire_tableaux_pdf(fichier_pdf)

# Créer le dossier de sortie si nécessaire
os.makedirs(dossier_sortie, exist_ok=True)

# Créer le document Word avec les données extraites
creer_document_word(texte, images, liens, tableaux, chemin_fichier_sortie)

print(f"Extraction terminée pour {fichier_pdf}. Résultats écrits dans : {chemin_fichier_sortie}")
